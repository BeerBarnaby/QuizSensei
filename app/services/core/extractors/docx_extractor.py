"""
DOCX text extraction strategy.
Uses python-docx to parse and extract structured text from Microsoft Word documents.
"""
import asyncio
import io
import base64
import logging
from pathlib import Path
import docx

from app.core.config import get_settings
from app.services.extractors.base import BaseExtractor

class DocxExtractor(BaseExtractor):
    """Extracts text from .docx documents."""

    def __init__(self):
        self.settings = get_settings()

    def _extract_sync(self, path: Path) -> str:
        """Synchronous wrapper around python-docx logic."""
        document = docx.Document(str(path))

        # Extract paragraph text
        text_parts = []
        for para in document.paragraphs:
            if para.text:
                text_parts.append(para.text)

        # Also extract table text as it's often missing if only paragraphs are used
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        return "\n".join(text_parts).strip()

    async def extract_text(self, file_path: Path) -> str:
        """
        Asynchronously extract text from a DOCX file.
        """
        loop = asyncio.get_running_loop()
        text = await loop.run_in_executor(None, self._extract_sync, file_path)
        return text
