"""
Unified extraction service for QuizSensei.
Supports PDF, DOCX, and TXT files using a factory pattern.
"""
import logging
import asyncio
import re
import io
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Dict, Type

import aiofiles
import docx
from pypdf import PdfReader

from app.core.config import get_settings

logger = logging.getLogger(__name__)

class BaseExtractor(ABC):
    @abstractmethod
    async def extract(self, file_path: Path) -> str:
        pass

class PDFExtractor(BaseExtractor):
    def _clean_text(self, text: str) -> str:
        if not text: return ""
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            stripped = line.strip()
            if re.match(r'^[\s\-\._=…]{5,}$', stripped): continue
            line = re.sub(r'[\-\._=…]{5,}', ' ', line)
            cleaned_lines.append(line.rstrip())
        text = '\n'.join(cleaned_lines)
        return re.sub(r'\n{3,}', '\n\n', text).strip()

    async def extract(self, file_path: Path) -> str:
        try:
            def _read():
                reader = PdfReader(file_path)
                return "\n\n".join(p.extract_text() or "" for p in reader.pages).strip()
            raw = await asyncio.to_thread(_read)
            return self._clean_text(raw)
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return ""

class DocxExtractor(BaseExtractor):
    def _read_sync(self, path: Path) -> str:
        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs if p.text]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip(): parts.append(cell.text)
        return "\n".join(parts).strip()

    async def extract(self, file_path: Path) -> str:
        return await asyncio.get_running_loop().run_in_executor(None, self._read_sync, file_path)

class TxtExtractor(BaseExtractor):
    async def extract(self, file_path: Path) -> str:
        async with aiofiles.open(file_path, mode="r", encoding="utf-8") as f:
            return await f.read()

class ExtractionService:
    """Factory service for document text extraction."""
    
    _EXTRACTORS: Dict[str, Type[BaseExtractor]] = {
        ".pdf": PDFExtractor,
        ".docx": DocxExtractor,
        ".txt": TxtExtractor
    }

    @classmethod
    async def extract_text(cls, file_path: Path) -> str:
        ext = file_path.suffix.lower()
        extractor_cls = cls._EXTRACTORS.get(ext)
        if not extractor_cls:
            raise ValueError(f"Unsupported file extension: {ext}")
        
        extractor = extractor_cls()
        return await extractor.extract(file_path)
